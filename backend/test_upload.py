import requests
from docx import Document

# 1. Create a dummy test resume (.docx)
print("Creating a dummy test resume...")
doc = Document()
doc.add_paragraph("This is a test resume for John Doe.")
doc.add_paragraph("Skills: python, software engineering, machine learning, deep learning, artificial intelligence.")
doc.save("test_resume.docx")

# 2. Send the file to your FastAPI backend
url = "http://127.0.0.1:8000/analyze"
print(f"Sending test file to {url}...")

with open("test_resume.docx", "rb") as f:
    files = {"file": ("test_resume.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    response = requests.post(url, files=files)

# 3. Print the result!
print("="*40)
print(f"Status Code: {response.status_code}")
print("="*40)
print("Response from Backend (including Supabase status):")
import pprint
pprint.pprint(response.json())
